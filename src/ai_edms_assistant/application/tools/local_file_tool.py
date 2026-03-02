# src/ai_edms_assistant/application/tools/local_file_tool.py
"""LocalFileTool — extracts text from locally uploaded files.
"""
from __future__ import annotations

import logging
import zipfile
from pathlib import Path
from typing import Any

import structlog
from pydantic import BaseModel, Field

from ...domain.value_objects.file_metadata import FileMetadata
from ..processors import AbstractProcessor
from .base_tool import AbstractEdmsTool

log = structlog.get_logger(__name__)
logger = logging.getLogger(__name__)

_SUPPORTED_EXTENSIONS: tuple[str, ...] = (
    ".pdf",
    ".docx",
    ".doc",
    ".txt",
    ".rtf",
    ".html",
    ".htm",
)


class LocalFileInput(BaseModel):
    """Input schema for local file reading.

    Attributes:
        file_path: Full absolute path to local file.
    """

    file_path: str = Field(
        ...,
        description=(
            "ПОЛНЫЙ абсолютный путь к локальному файлу. "
            "Возьми его ИЗ ТЕКСТА ЗАПРОСА в блоке [ДОСТУПЕН ЛОКАЛЬНЫЙ ФАЙЛ] "
            "или из поля file_path контекста."
        ),
    )


class LocalFileTool(AbstractEdmsTool):
    """Tool for extracting text from locally uploaded files.

    Supports: PDF, DOCX, TXT, DOC, RTF via FileProcessor or built-in fallback.

    Extraction strategy for DOCX (in order of preference):
        1. python-docx: paragraphs + table cells + text boxes.
        2. XML fallback: raw w:t element text from document.xml via zipfile.
        3. If both return empty — returns descriptive error, NOT empty string.

    This prevents the agent validator from misreading an empty result
    as a tool failure (which caused LLM to respond "не удалось прочитать файл").
    """

    name: str = "read_local_file_content"
    description: str = (
        "Извлекает текст из локального файла (PDF, DOCX, TXT). "
        "Используй ТОЛЬКО для файлов загруженных пользователем напрямую через интерфейс. "
        "НЕ используй для вложений EDMS — для них есть doc_get_file_content."
    )
    args_schema: type[BaseModel] = LocalFileInput

    def __init__(
        self,
        file_processor: AbstractProcessor[FileMetadata, str] | None = None,
        **kwargs: Any,
    ) -> None:
        """Initialize with optional file processor.

        Args:
            file_processor: Processor for text extraction.
                If None, uses built-in fallback logic.
            **kwargs: Additional BaseTool arguments.
        """
        super().__init__(**kwargs)
        self._processor = file_processor

    async def _arun(self, file_path: str, **kwargs: Any) -> dict[str, Any]:
        """Execute local file text extraction.

        Args:
            file_path: Absolute path to local file. Injected by
                ToolParameterInjector from context — never trust LLM value alone.

        Returns:
            Success dict with extracted text and metadata,
            or error dict when extraction fails.
        """
        try:
            # ── Guard: некорректный аргумент от LLM ──────────────────────────
            cleaned = str(file_path).strip()
            if not cleaned or cleaned.lower() in ("file_path", "path", "none", "null"):
                return self._handle_error(
                    ValueError(
                        "Передан некорректный путь. "
                        "Инструмент требует абсолютный путь к файлу."
                    )
                )

            path = Path(cleaned)

            if not path.exists():
                log.warning(
                    "local_file_not_found",
                    path=cleaned,
                    path_tail=cleaned[-60:],
                )
                return self._handle_error(
                    ValueError(f"Файл не найден: {path.name}")
                )

            ext = path.suffix.lower()
            if ext not in _SUPPORTED_EXTENSIONS:
                return self._handle_error(
                    ValueError(
                        f"Формат {ext} не поддерживается. "
                        f"Поддерживаются: {', '.join(_SUPPORTED_EXTENSIONS)}"
                    )
                )

            # ── Extract text ──────────────────────────────────────────────────
            if self._processor:
                file_meta = self._build_file_metadata(path)
                text_content = await self._processor.process(file_meta)
            else:
                text_content = await self._extract_text(path)

            # ── Empty result guard ────────────────────────────────────────────
            if not text_content or not text_content.strip():
                log.warning(
                    "local_file_empty_text",
                    file_name=path.name,
                    size_bytes=path.stat().st_size,
                )
                return self._handle_error(
                    ValueError(
                        f"Файл «{path.name}» не содержит извлекаемого текста. "
                        f"Возможно, документ содержит только изображения или защищён от копирования."
                    )
                )

            log.debug(
                "text_extracted_via_shared_utils",
                filename=path.name,
                chars=len(text_content),
            )

            preview = text_content[:15_000]
            is_truncated = len(text_content) > 15_000

            return self._success_response(
                data={
                    "meta": {
                        "name": path.name,
                        "size_kb": round(path.stat().st_size / 1024, 1),
                        "extension": ext.lstrip("."),
                    },
                    "content": preview,
                    "is_truncated": is_truncated,
                    "total_chars": len(text_content),
                },
                message="Текст извлечён",
            )

        except Exception as exc:
            log.error(
                "local_file_extraction_failed",
                file_path=str(file_path)[-60:],
                error=str(exc),
            )
            return self._handle_error(exc)

    # ── Text extraction ───────────────────────────────────────────────────────

    async def _extract_text(self, path: Path) -> str:
        """Route extraction by file extension.

        Args:
            path: Validated path to existing file.

        Returns:
            Extracted text string (may be empty if document has no text).
        """
        ext = path.suffix.lower()
        if ext == ".txt":
            return self._read_txt(path)
        if ext in (".html", ".htm"):
            return self._read_html(path)
        if ext == ".pdf":
            return self._read_pdf(path)
        if ext == ".docx":
            return self._read_docx(path)
        if ext in (".doc", ".rtf"):
            return f"[Формат {ext} требует LibreOffice. Файл: {path.name}]"
        return ""

    # ── Format-specific readers ───────────────────────────────────────────────

    def _read_txt(self, path: Path) -> str:
        """Read plain text file.

        Args:
            path: Path to .txt file.

        Returns:
            File contents as string.
        """
        return path.read_text(encoding="utf-8", errors="ignore")

    def _read_html(self, path: Path) -> str:
        """Extract text from HTML, stripping tags.

        Args:
            path: Path to .html / .htm file.

        Returns:
            Stripped text content.
        """
        content = path.read_text(encoding="utf-8", errors="ignore")
        try:
            from bs4 import BeautifulSoup  # type: ignore[import]
            return BeautifulSoup(content, "html.parser").get_text(
                separator="\n", strip=True
            )
        except ImportError:
            import re
            content = re.sub(r"<[^>]+>", " ", content)
            return re.sub(r"\s+", " ", content).strip()

    def _read_pdf(self, path: Path) -> str:
        """Extract text from PDF using pypdf (preferred) or PyPDF2 fallback.

        Args:
            path: Path to .pdf file.

        Returns:
            Concatenated page text.
        """
        try:
            import pypdf  # type: ignore[import]
            parts: list[str] = []
            with open(path, "rb") as f:
                reader = pypdf.PdfReader(f)
                for page in reader.pages:
                    txt = page.extract_text()
                    if txt:
                        parts.append(txt)
            return "\n".join(parts)
        except ImportError:
            pass

        try:
            import PyPDF2
            parts = []
            with open(path, "rb") as f:
                reader = PyPDF2.PdfReader(f)
                for page in reader.pages:
                    txt = page.extract_text()
                    if txt:
                        parts.append(txt)
            return "\n".join(parts)
        except ImportError:
            return (
                f"[PDF парсинг недоступен. Установите: pip install pypdf]\n"
                f"Файл: {path.name}"
            )

    def _read_docx(self, path: Path) -> str:
        """Extract all text from DOCX: paragraphs + tables + shapes + XML fallback.

        Strategy:
            1. python-docx: iterate paragraphs and table cells.
            2. TextBoxes via XML namespace (often missed by python-docx high-level API).
            3. zipfile XML fallback: extract all <w:t> elements from document.xml.
            4. If all empty — return empty string (caller handles this case).

        Args:
            path: Path to .docx file.

        Returns:
            Extracted text (may be empty for image-only documents).
        """
        try:
            import docx as _docx
            from lxml import etree

            doc = _docx.Document(str(path))
            parts: list[str] = []

            for para in doc.paragraphs:
                t = para.text.strip()
                if t:
                    parts.append(t)

            for table in doc.tables:
                for row in table.rows:
                    row_cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                    if row_cells:
                        parts.append(" | ".join(row_cells))

            nsmap = {
                "w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main",
                "mc": "http://schemas.openxmlformats.org/markup-compatibility/2006",
            }
            try:
                body = doc.element.body
                for txbx in body.iter(
                    "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}txbxContent"
                ):
                    for para_el in txbx.iter(
                        "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}p"
                    ):
                        texts = para_el.findall(
                            ".//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}t"
                        )
                        line = "".join(t.text or "" for t in texts).strip()
                        if line:
                            parts.append(f"[TextBox] {line}")
            except Exception:
                pass

            if parts:
                result = "\n".join(parts)
                log.debug(
                    "docx_extracted_via_python_docx",
                    filename=path.name,
                    chars=len(result),
                    paragraphs=len([p for p in doc.paragraphs if p.text.strip()]),
                    tables=len(doc.tables),
                )
                return result

        except ImportError:
            log.warning("python_docx_not_installed", file_name=path.name)
        except Exception as exc:
            log.warning(
                "python_docx_extraction_failed",
                file_name=path.name,
                error=str(exc),
            )

        return self._read_docx_xml_fallback(path)

    def _read_docx_xml_fallback(self, path: Path) -> str:
        """Extract text from DOCX by parsing document.xml directly via zipfile.

        Does not require python-docx. Works with any valid DOCX (ZIP) file.
        Extracts all <w:t> element text in document order.

        Args:
            path: Path to .docx file.

        Returns:
            Concatenated text from all w:t elements, or empty string.
        """
        try:
            import xml.etree.ElementTree as ET

            W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"

            with zipfile.ZipFile(str(path), "r") as zf:
                if "word/document.xml" not in zf.namelist():
                    return ""
                xml_bytes = zf.read("word/document.xml")

            root = ET.fromstring(xml_bytes)
            texts: list[str] = []

            for elem in root.iter(f"{{{W_NS}}}t"):
                t = (elem.text or "").strip()
                if t:
                    texts.append(t)

            result = " ".join(texts)
            log.debug(
                "docx_extracted_via_xml_fallback",
                filename=path.name,
                chars=len(result),
                w_t_count=len(texts),
            )
            return result

        except zipfile.BadZipFile:
            log.warning("docx_bad_zip", file_name=path.name)
            return ""
        except Exception as exc:
            log.warning(
                "docx_xml_fallback_failed",
                file_name=path.name,
                error=str(exc),
            )
            return ""

    def _build_file_metadata(self, path: Path) -> FileMetadata:
        """Build FileMetadata value object from path.

        Args:
            path: Path to file.

        Returns:
            FileMetadata value object.
        """
        from uuid import uuid4
        from ...domain.entities.attachment import ContentType

        ext = path.suffix.lower()
        content_type_map: dict[str, ContentType] = {
            ".pdf": ContentType.PDF,
            ".docx": ContentType.DOCX,
            ".doc": ContentType.DOC,
            ".txt": ContentType.TXT,
            ".rtf": ContentType.RTF,
            ".html": ContentType.HTML,
            ".htm": ContentType.HTML,
        }
        content_type = content_type_map.get(ext, ContentType.UNKNOWN)

        return FileMetadata(
            attachment_id=uuid4(),
            file_name=path.name,
            content_type=content_type,
            size_bytes=path.stat().st_size,
            storage_path=str(path.absolute()),
        )

    def _run(self, *args: Any, **kwargs: Any) -> None:
        """Sync execution not supported.

        Raises:
            NotImplementedError: Always.
        """
        raise NotImplementedError("Use _arun for async execution")